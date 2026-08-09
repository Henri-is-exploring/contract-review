#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""generate_report.py — 由 JSON 审核结果生成 .docx 审核报告。

用法:
  $PY generate_report.py --input review_result.json --output review_report.docx

输入 JSON 格式:
  {
    "contract_name": "货物买卖合同",
    "review_date": "2026-08-09",
    "our_role": "甲方（买方）",
    "negotiation_position": "弱势方",
    "overall_risk": {"level": "中", "reason": "..."},
    "clauses": [
      {"priority":"高","name":"标的物条款","risk":"...","suggestion":"...","reason":"..."},
      {"priority":"中","name":"交付条款","risk":"...","suggestion":"...","reason":"..."}
    ],
    "todos": ["确认...","确认..."]
  }

报告三部分(缺一算失败):
  1. 整体风险等级(高/中/低 + 理由)
  2. 条款建议按高/中/低优先级分组,每条:条款名/风险说明/修改建议/理由
  3. TODO: 需法务向业务确认的信息
"""
import argparse
import json
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RISK_COLOR = {'高': RGBColor(0xC0, 0x00, 0x00),
              '中': RGBColor(0xBF, 0x80, 0x00),
              '低': RGBColor(0x00, 0x70, 0x30)}
PRIORITY_ORDER = ['高', '中', '低']


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def _kv(doc, key, val):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{key}：')
    r1.bold = True
    p.add_run(val)
    return p


def build(report, out):
    doc = Document()

    # 标题
    title = doc.add_heading('合同审核报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    _kv(doc, '合同名称', report.get('contract_name', '—'))
    _kv(doc, '审核日期', report.get('review_date', datetime.now().strftime('%Y-%m-%d')))
    _kv(doc, '我方立场', report.get('our_role', '—'))
    _kv(doc, '谈判地位', report.get('negotiation_position', '—'))
    doc.add_paragraph('')

    # ===== 一、整体风险等级 =====
    risk = report.get('overall_risk', {})
    level = risk.get('level', '未评定')
    _heading(doc, '一、整体风险等级', level=1)
    _para(doc, f'风险等级：{level}', bold=True,
          color=RISK_COLOR.get(level))
    _kv(doc, '评定理由', risk.get('reason', '—'))
    doc.add_paragraph('')

    # ===== 二、条款修改建议（按优先级分组）=====
    _heading(doc, '二、条款修改建议', level=1)
    clauses = report.get('clauses', [])
    by_pri = {p: [c for c in clauses if c.get('priority') == p] for p in PRIORITY_ORDER}
    other = [c for c in clauses if c.get('priority') not in PRIORITY_ORDER]
    if other:
        by_pri['未分级'] = other

    found_any = False
    for pri in list(by_pri.keys()):
        items = by_pri[pri]
        if not items:
            continue
        found_any = True
        _heading(doc, f'（{PRIORITY_ORDER.index(pri)+1 if pri in PRIORITY_ORDER else "·"}）{pri}优先级', level=2)
        for i, c in enumerate(items, 1):
            _kv(doc, '条款名称', c.get('name', '—'))
            _kv(doc, '风险说明', c.get('risk', '—'))
            _kv(doc, '修改建议', c.get('suggestion', '—'))
            _kv(doc, '理由', c.get('reason', '—'))
            doc.add_paragraph('')
    if not found_any:
        _para(doc, '暂无条款修改建议。')
    doc.add_paragraph('')

    # ===== 三、待确认事项（TODO）=====
    _heading(doc, '三、待确认事项（需向业务确认）', level=1)
    todos = report.get('todos', [])
    if not todos:
        _para(doc, '暂无待确认事项。')
    else:
        for i, t in enumerate(todos, 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(t)

    doc.save(out)


def main():
    ap = argparse.ArgumentParser(description='由 JSON 审核结果生成 .docx 审核报告')
    ap.add_argument('--input', required=True, help='审核结果 JSON 路径')
    ap.add_argument('--output', required=True, help='输出 .docx 路径')
    args = ap.parse_args()
    with open(args.input, encoding='utf-8') as f:
        report = json.load(f)
    build(report, args.output)
    n = len(report.get('clauses', []))
    print(f'OK -> {args.output} | 风险:{report.get("overall_risk",{}).get("level","?")} '
          f'条款:{n} TODO:{len(report.get("todos",[]))}', file=sys.stderr)


if __name__ == '__main__':
    main()
