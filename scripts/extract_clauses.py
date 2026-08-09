#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""extract_clauses.py — 从 .docx 合同中抽取条款清单（条号 + 主旨 + 原文片段）。

用途：合同审核第 4.0 步强制前置，先把合同所有条款抽成清单，再逐条对照
playbook 判断是否命中，防止 LLM 在长合同里跳过不显眼条款。

用法:
  $PY extract_clauses.py --input 原合同.docx --output clauses.json

输出 JSON:
  [
    {"para_idx":5,"level":"article","marker":"第一条","heading":"第一条  合作模式与订单机制","snippet":""},
    {"para_idx":6,"level":"clause","marker":"1.1","heading":"","snippet":"本协议为年度框架采购协议……"},
    {"para_idx":7,"level":"clause","marker":"1.2","heading":"","snippet":"乙方收到甲方订单后……"}
  ]

说明:
  - para_idx 与 python-docx 的 Document(path).paragraphs 索引对齐（0 起，含标题段），
    便于后续 track_changes.py 的 para_idx 引用。
  - level=article：以「第X条」起头的段落；level=clause：以「数字.数字」起头的段落。
  - snippet 取段落纯文本前 120 字。
  - 既非 article 也非 clause 的段落不输出（视为正文叙述/签署栏，不在条款清单内）。
"""
import argparse
import json
import re
import sys

from docx import Document

ARTICLE_RE = re.compile(r'^(第[一二三四五六七八九十百千零]+条)')
CLAUSE_RE = re.compile(r'^(\d+\.\d+)')
SNIPPET_LEN = 120


def extract(path):
    doc = Document(path)
    out = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        m_art = ARTICLE_RE.match(text)
        if m_art:
            out.append({
                "para_idx": idx,
                "level": "article",
                "marker": m_art.group(1),
                "heading": text,
                "snippet": "",
            })
            continue
        m_cl = CLAUSE_RE.match(text)
        if m_cl:
            out.append({
                "para_idx": idx,
                "level": "clause",
                "marker": m_cl.group(1),
                "heading": "",
                "snippet": text[:SNIPPET_LEN],
            })
            continue
    return out


def main():
    ap = argparse.ArgumentParser(description='从 .docx 合同抽取条款清单')
    ap.add_argument('--input', required=True, help='合同 .docx 路径')
    ap.add_argument('--output', required=True, help='输出 clauses.json 路径')
    args = ap.parse_args()
    clauses = extract(args.input)
    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(clauses, f, ensure_ascii=False, indent=2)
    n_art = sum(1 for c in clauses if c['level'] == 'article')
    n_cl = sum(1 for c in clauses if c['level'] == 'clause')
    print(f'OK -> {args.output} | 条 {n_art} / 款 {n_cl} / 合计 {len(clauses)}', file=sys.stderr)


if __name__ == '__main__':
    main()
